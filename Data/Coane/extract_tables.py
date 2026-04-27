"""
Extract tables A1 and A2 from the supplementary materials docx file into CSV files.

Usage:
    python extract_tables.py [input.docx] [--output-dir DIR]

Defaults:
    input  : 13428_2020_1493_MOESM1_ESM.docx  (same directory as the script)
    output : same directory as the script
"""

import csv
import sys
import argparse
from pathlib import Path
from docx import Document


def get_merged_headers(table):
    """
    Build a two-row merged header for a table.

    Row 0 contains spanning group labels; row 1 contains per-column labels.
    Returns a list of final column header strings.
    """
    # Collect raw cell text for first two rows
    row0 = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
    row1 = [cell.text.strip().replace("\n", " ") for cell in table.rows[1].cells]

    headers = []
    prev_group = None
    for i, (group, col) in enumerate(zip(row0, row1)):
        # Blank group cell means "same group as previous column"
        effective_group = group if group else prev_group
        prev_group = effective_group

        if effective_group and col and effective_group.lower() != col.lower():
            headers.append(f"{effective_group} – {col}")
        elif col:
            headers.append(col)
        elif effective_group:
            headers.append(effective_group)
        else:
            headers.append(f"Column_{i+1}")

    return headers


def extract_table(table, table_name):
    """
    Extract all data rows from a table, returning (headers, rows).

    Assumes the first two rows are header rows. Word repeats those header rows
    (preceded by blank spacer rows) at every page break — these are real table
    rows, not a python-docx artefact. We detect and skip them by comparing each
    row against the two canonical header signatures.
    """
    headers = get_merged_headers(table)

    # Build normalised signatures for both header rows so we can detect repeats
    def normalise(row):
        return tuple(c.text.strip().replace("\xa0", "").replace("\n", " ")
                     for c in row.cells)

    header_sig0 = normalise(table.rows[0])
    header_sig1 = normalise(table.rows[1])

    rows = []
    skipped = 0
    for row in table.rows[2:]:          # skip the original two header rows
        cells = [
            c.text.strip().replace("\xa0", "").replace("\n", " ")
            for c in row.cells
        ]
        sig = tuple(cells)

        # Skip blank rows and any row that matches either header signature
        if not any(cells):
            continue
        if sig == header_sig0 or sig == header_sig1:
            skipped += 1
            continue

        rows.append(cells)

    print(f"  {table_name}: {len(rows)} data rows, {len(headers)} columns "
          f"(skipped {skipped} repeated header rows)")
    return headers, rows


def save_csv(headers, rows, path):
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        writer.writerows(rows)
    print(f"  Saved → {path}")


def main():
    parser = argparse.ArgumentParser(description="Extract tables A1 & A2 to CSV")
    parser.add_argument(
        "input",
        nargs="?",
        default="13428_2020_1493_MOESM1_ESM.docx",
        help="Path to the input .docx file",
    )
    parser.add_argument(
        "--output-dir",
        default=None,
        help="Directory for output CSVs (defaults to same dir as input file)",
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        sys.exit(f"Error: file not found: {input_path}")

    output_dir = Path(args.output_dir) if args.output_dir else input_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Reading: {input_path}")
    doc = Document(input_path)

    if len(doc.tables) < 2:
        sys.exit(f"Error: expected 2 tables, found {len(doc.tables)}")

    table_names = ["A1", "A2"]
    for i, (table, name) in enumerate(zip(doc.tables, table_names)):
        print(f"\nExtracting Table {name}:")
        headers, rows = extract_table(table, name)
        out_path = output_dir / f"table_{name}.csv"
        save_csv(headers, rows, out_path)

    print("\nDone.")


if __name__ == "__main__":
    main()